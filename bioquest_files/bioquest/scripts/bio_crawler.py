"""
BioQuest 高中生物题目爬虫 v1.1

目标：从允许爬取的免费公开网站抓取高中生物题目（含解析），
输出格式与 server.py / pool.json 保持一致，可直接合并入库或同步到 Supabase。

合规约定：
- 爬取前检查 robots.txt
- 默认限速 2 次/秒，可通过 --rate 调整
- 仅抓取公开页面，不绕过登录、不破解付费内容

依赖：
    pip install requests beautifulsoup4 lxml pymupdf

用法示例：
    # 1. 爬高考题（HTML 页面），作为高考模式练习题并同步到 Supabase
    python scripts/bio_crawler.py --mode html \
        --start-urls "https://example.com/gaokao/bio" \
        --target high_school --difficulty basic \
        --max-pages 3 --sync

    # 2. 爬竞赛题（PDF 真题），作为竞赛模式练习题
    python scripts/bio_crawler.py --mode pdf \
        --pdf-urls "https://example.com/league/2024bio.pdf" \
        --target competition --difficulty league \
        --sync

    # 3. 把结果追加到 pool.json
    python scripts/bio_crawler.py --merge-to pool.json

概念标签已与 js/knowledge-graph.js 对齐，确保图谱点击节点可进入专项练习。
"""

import argparse
import json
import logging
import os
import random
import re
import sys
import time
from pathlib import Path
from urllib.parse import urljoin, urlparse
from urllib.robotparser import RobotFileParser

import requests
from bs4 import BeautifulSoup

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)
log = logging.getLogger("bio_crawler")

# 复用 server.py 的 Supabase 工具（credentials、sb_request、_build_record 等）
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)
try:
    from server import sb_request, _build_record, sb_upsert_batch
    HAS_SERVER = True
    log.info("已复用 server.py 的 Supabase 配置与函数")
except Exception as e:
    HAS_SERVER = False
    log.warning(f"无法导入 server.py，Supabase 同步不可用: {e}")



# ═══════════════════════════════════════════════════════════════
# 数据模型：与 server.py pool.json 格式保持一致
# ═══════════════════════════════════════════════════════════════

# ═══════════════════════════════════════════════════════════════
# 知识图谱对齐：概念、学科、模块必须与 js/knowledge-graph.js 一致
# ═══════════════════════════════════════════════════════════════

# concept -> (category, module)
# label 必须和 knowledge-graph.js 中 GRAPH_NODES 的 label 完全一致
CONCEPT_MAP = {
    # 细胞生物学 (module_1)
    "细胞结构": ("细胞生物学", "module_1"),
    "细胞膜": ("细胞生物学", "module_1"),
    "细胞器": ("细胞生物学", "module_1"),
    "细胞周期": ("细胞生物学", "module_1"),
    "细胞信号转导": ("细胞生物学", "module_1"),
    "细胞凋亡": ("细胞生物学", "module_1"),
    # 遗传学 (module_4)
    "孟德尔遗传": ("遗传学", "module_4"),
    "连锁与交换": ("遗传学", "module_4"),
    "伴性遗传": ("遗传学", "module_4"),
    "基因突变": ("遗传学", "module_4"),
    "染色体变异": ("遗传学", "module_4"),
    "群体遗传": ("遗传学", "module_4"),
    # 分子生物学 (module_1)
    "DNA结构": ("分子生物学", "module_1"),
    "DNA复制": ("分子生物学", "module_1"),
    "转录": ("分子生物学", "module_1"),
    "翻译": ("分子生物学", "module_1"),
    "基因表达调控": ("分子生物学", "module_1"),
    "PCR技术": ("分子生物学", "module_1"),
    # 生态学 (module_3)
    "生态系统": ("生态学", "module_3"),
    "种群生态": ("生态学", "module_3"),
    "群落生态": ("生态学", "module_3"),
    "物质循环": ("生态学", "module_3"),
    "生物多样性": ("生态学", "module_3"),
    # 动物学 (module_3)
    "动物组织": ("动物学", "module_3"),
    "神经系统": ("动物学", "module_3"),
    "免疫系统": ("动物学", "module_3"),
    "内分泌系统": ("动物学", "module_3"),
    "循环系统": ("动物学", "module_3"),
    # 植物学 (module_2)
    "植物组织": ("植物学", "module_2"),
    "光合作用": ("植物学", "module_2"),
    "植物激素": ("植物学", "module_2"),
    "植物物质运输": ("植物学", "module_2"),
    # 生物化学 (module_1)
    "酶": ("生物化学", "module_1"),
    "糖酵解": ("生物化学", "module_1"),
    "柠檬酸循环": ("生物化学", "module_1"),
    "氧化磷酸化": ("生物化学", "module_1"),
    "氨基酸代谢": ("生物化学", "module_1"),
    # 微生物学 (module_2)
    "细菌": ("微生物学", "module_2"),
    "病毒": ("微生物学", "module_2"),
    "微生物生态": ("微生物学", "module_2"),
}

# 为每个概念定义触发关键词（用于从题干/解析中匹配）
CONCEPT_KEYWORDS = {
    "细胞结构": ["细胞结构", "原核细胞", "真核细胞", "细胞质", "细胞核"],
    "细胞膜": ["细胞膜", "流动镶嵌", "跨膜运输", "被动运输", "主动运输"],
    "细胞器": ["细胞器", "线粒体", "叶绿体", "内质网", "高尔基体", "溶酶体", "核糖体"],
    "细胞周期": ["细胞周期", "有丝分裂", "减数分裂", "分裂间期", "分裂期"],
    "细胞信号转导": ["信号转导", "第二信使", "G蛋白", "受体", "信号通路"],
    "细胞凋亡": ["细胞凋亡", "程序性死亡", "Caspase"],
    "孟德尔遗传": ["孟德尔", "分离定律", "自由组合", "显隐性"],
    "连锁与交换": ["连锁", "交换", "重组率", "三点测交"],
    "伴性遗传": ["伴性遗传", "X连锁", "Y连锁", "巴氏小体"],
    "基因突变": ["基因突变", "碱基替换", "移码突变"],
    "染色体变异": ["染色体变异", "缺失", "重复", "倒位", "易位", "非整倍体"],
    "群体遗传": ["群体遗传", "Hardy-Weinberg", "遗传漂变", "基因流"],
    "DNA结构": ["DNA结构", "双螺旋", "Chargaff", "碱基配对"],
    "DNA复制": ["DNA复制", "半保留复制", "复制叉", "冈崎片段"],
    "转录": ["转录", "RNA聚合酶", "启动子"],
    "翻译": ["翻译", "核糖体", "密码子", "tRNA"],
    "基因表达调控": ["基因表达调控", "操纵子", "转录因子", "表观遗传"],
    "PCR技术": ["PCR", "聚合酶链式反应", "引物设计"],
    "生态系统": ["生态系统", "食物链", "食物网", "能量流动"],
    "种群生态": ["种群", "逻辑斯谛", "r选择", "K选择"],
    "群落生态": ["群落", "种间关系", "演替", "生态位"],
    "物质循环": ["物质循环", "碳循环", "氮循环", "磷循环"],
    "生物多样性": ["生物多样性", "物种多样性", "遗传多样性"],
    "动物组织": ["动物组织", "上皮组织", "结缔组织", "肌肉组织", "神经组织"],
    "神经系统": ["神经系统", "神经元", "突触", "反射弧"],
    "免疫系统": ["免疫系统", "抗体", "T细胞", "B细胞", "免疫应答"],
    "内分泌系统": ["内分泌", "激素", "垂体", "甲状腺"],
    "循环系统": ["循环系统", "心脏", "血液循环", "血压"],
    "植物组织": ["植物组织", "分生组织", "输导组织", "薄壁组织"],
    "光合作用": ["光合作用", "光反应", "暗反应", "Calvin循环"],
    "植物激素": ["植物激素", "生长素", "赤霉素", "细胞分裂素", "乙烯", "脱落酸"],
    "植物物质运输": ["植物物质运输", "蒸腾作用", "筛管", "导管"],
    "酶": ["酶", "米氏方程", "竞争性抑制", "别构调节"],
    "糖酵解": ["糖酵解", "丙酮酸", "底物水平磷酸化"],
    "柠檬酸循环": ["柠檬酸循环", "三羧酸循环", "Krebs循环"],
    "氧化磷酸化": ["氧化磷酸化", "电子传递链", "化学渗透", "ATP合酶"],
    "氨基酸代谢": ["氨基酸代谢", "转氨基", "脱氨基", "尿素循环"],
    "细菌": ["细菌", "革兰氏染色", "细菌形态"],
    "病毒": ["病毒", "病毒复制", "溶原性", "溶菌性"],
    "微生物生态": ["微生物生态", "共生", "寄生", "极端微生物"],
}

CATEGORY_MODULE = {
    "细胞生物学": "module_1",
    "分子生物学": "module_1",
    "生物化学": "module_1",
    "遗传学": "module_4",
    "植物学": "module_2",
    "微生物学": "module_2",
    "动物学": "module_3",
    "生态学": "module_3",
}


def infer_concept_subject_module(stem: str, analysis: str = "") -> tuple:
    """
    返回 (concept, subject, module)，全部与知识图谱对齐。
    优先按 CONCEPT_KEYWORDS 命中（题干权重 3 倍，解析权重 1 倍）；
    未命中则按学科关键词推断。
    """
    stem = stem or ""
    analysis = analysis or ""
    best_concept, best_score = "", 0
    for concept, kws in CONCEPT_KEYWORDS.items():
        score_stem = sum(stem.count(k) for k in kws) * 3
        score_ana = sum(analysis.count(k) for k in kws)
        score = score_stem + score_ana
        if score > best_score:
            best_score = score
            best_concept = concept

    if best_concept:
        subject, module = CONCEPT_MAP[best_concept]
        return best_concept, subject, module

    # 兜底：按学科关键词（同样题干加权）
    subject_scores = {}
    for subj, kws in {
        "细胞生物学": ["细胞", "有丝分裂", "减数分裂"],
        "遗传学": ["遗传", "孟德尔", "伴性", "基因突变", "染色体"],
        "分子生物学": ["DNA", "RNA", "复制", "转录", "翻译", "基因表达"],
        "生态学": ["生态", "种群", "群落", "演替", "生态系统"],
        "动物学": ["动物", "神经", "免疫", "内分泌", "循环"],
        "植物学": ["植物", "激素", "光合"],
        "生物化学": ["酶", "代谢", "呼吸", "糖酵解"],
        "微生物学": ["细菌", "病毒", "微生物"],
    }.items():
        subject_scores[subj] = sum(stem.count(k) * 3 for k in kws) + sum(analysis.count(k) for k in kws)

    best_subject = max(subject_scores, key=subject_scores.get)
    if subject_scores[best_subject] == 0:
        best_subject = "生物学"
    module = CATEGORY_MODULE.get(best_subject, "module_1")
    return "", best_subject, module


def build_question(stem: str, options: dict, answer, analysis: str = "",
                   year: int = None, source: str = "", target: str = "high_school",
                   difficulty: str = "medium") -> dict:
    """构造与 pool.json 兼容的题目记录。"""
    concept, subject, module = infer_concept_subject_module(stem, analysis)

    # answer 兼容单字母字符串或 dict（支持 A-E）
    all_opts = ["A", "B", "C", "D", "E"]
    if isinstance(answer, str) and len(answer) == 1 and answer in "ABCDE":
        answer_obj = {k: (k == answer) for k in all_opts}
    elif isinstance(answer, dict):
        answer_obj = {k: bool(answer.get(k, False)) for k in all_opts}
    else:
        answer_obj = {k: False for k in all_opts}

    tags = [module, subject]
    if concept:
        tags.append(concept)
    tags.append(difficulty)
    tags.append(target)
    if year:
        tags.append(str(year))
    if source:
        tags.append(source)

    return {
        "id": f"crawled_{int(time.time()*1000)}_{random.randint(0,9999):04d}",
        "stem": stem.strip(),
        "options": {k: options.get(k, "") for k in all_opts},
        "answer": answer_obj,
        "analysis": analysis.strip(),
        "knowledge": [module, subject],
        "module": module,
        "difficulty": difficulty,
        "target": target,
        "subject": subject,
        "concept": concept,
        "tags": tags,
        "weight": 1.0,
        "fb_good": 0,
        "fb_bad": 0,
        "created_at": time.time(),
        "source": source,
        "year": year
    }


# ═══════════════════════════════════════════════════════════════
# 基础爬虫：限速 + robots.txt
# ═══════════════════════════════════════════════════════════════

class EthicalFetcher:
    def __init__(self, rate: float = 2.0, user_agent: str = None):
        self.delay = 1.0 / rate
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": user_agent or (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/120.0.0.0 Safari/537.36"
            ),
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "zh-CN,zh;q=0.9",
        })
        self.last_request = 0.0
        self.robots_cache = {}

    def _check_robots(self, url: str) -> bool:
        parsed = urlparse(url)
        robots_url = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
        if robots_url not in self.robots_cache:
            rp = RobotFileParser()
            rp.set_url(robots_url)
            try:
                rp.read()
            except Exception as e:
                log.warning(f"robots.txt 读取失败 {robots_url}: {e}")
            self.robots_cache[robots_url] = rp
        return self.robots_cache[robots_url].can_fetch("*", url)

    def get(self, url: str, **kwargs):
        if not self._check_robots(url):
            log.warning(f"[跳过] robots.txt 禁止: {url}")
            return None

        elapsed = time.time() - self.last_request
        if elapsed < self.delay:
            time.sleep(self.delay - elapsed)

        try:
            resp = self.session.get(url, timeout=20, **kwargs)
            self.last_request = time.time()
            resp.raise_for_status()
            return resp
        except requests.RequestException as e:
            log.error(f"[请求失败] {url}: {e}")
            return None

    def get_text(self, url: str) -> str:
        resp = self.get(url)
        return resp.text if resp else ""

    def get_rendered_html(self, url: str, wait_seconds: float = 2.0) -> str:
        """使用 Playwright 渲染 JS 后返回 HTML（用于 SPA/异步加载页面）。"""
        if not self._check_robots(url):
            log.warning(f"[跳过] robots.txt 禁止: {url}")
            return ""

        elapsed = time.time() - self.last_request
        if elapsed < self.delay:
            time.sleep(self.delay - elapsed)

        try:
            from playwright.sync_api import sync_playwright
            with sync_playwright() as p:
                browser = p.chromium.launch(headless=True)
                page = browser.new_page()
                page.goto(url, timeout=30000, wait_until="networkidle")
                page.wait_for_timeout(int(wait_seconds * 1000))
                html = page.content()
                browser.close()
            self.last_request = time.time()
            return html
        except Exception as e:
            log.error(f"[Playwright 渲染失败] {url}: {e}")
            return ""

    def download(self, url: str, save_path: str) -> bool:
        resp = self.get(url, stream=True)
        if not resp:
            return False
        try:
            with open(save_path, "wb") as f:
                for chunk in resp.iter_content(chunk_size=8192):
                    f.write(chunk)
            return True
        except Exception as e:
            log.error(f"[下载失败] {url}: {e}")
            return False


# ═══════════════════════════════════════════════════════════════
# 解析器注册表
# ═══════════════════════════════════════════════════════════════

class SiteParser:
    """站点解析器基类，子类需实现 can_handle / parse_list / parse_detail"""
    name = "base"
    target = "high_school"
    difficulty = "medium"

    def can_handle(self, url: str) -> bool:
        return False

    def parse_list(self, fetcher: EthicalFetcher, url: str, max_pages: int = 1) -> list:
        """返回详情页 URL 列表"""
        return []

    def parse_detail(self, fetcher: EthicalFetcher, url: str) -> list:
        """返回题目 dict 列表"""
        return []


PARSERS = []


def register_parser(parser_cls):
    PARSERS.append(parser_cls())
    return parser_cls


# ═══════════════════════════════════════════════════════════════
# 通用题目提取：基于正则/启发式，从任意文本中尝试提取选择题
# ═══════════════════════════════════════════════════════════════

class GenericQuestionExtractor:
    """不依赖特定 DOM 结构，从文本块中提取选择题。"""

    # 题号分隔：1. / 第1题 / (1) / 1．
    Q_SPLIT_RE = re.compile(
        r"(?:^|\n)\s*(?:第\s*(\d+)\s*题|(\d+)[\.．、]|\((\d+)\))\s*",
        re.MULTILINE
    )
    # 选项分隔：必须位于行首的 A. / A．/ A、/ A) / A）/ A（支持 A-E）
    OPT_RE = re.compile(
        r"(?:^|\n)\s*([A-E])[\.．、\)）]?\s+",
        re.MULTILINE
    )
    # 答案：答案：A / 【答案】B / 正确答案：C（支持 A-E）
    ANS_RE = re.compile(
        r"(?:答案|【答案】|正确答案)[：:\s为]+([A-E]+)",
        re.IGNORECASE
    )
    # 解析：解析：... / 【解析】...
    ANA_RE = re.compile(
        r"(?:解析|【解析】)[：:\s]*(.+?)(?=\n(?:第\s*\d+\s*题|\d+[\.．、]|\(\d+\))|$)",
        re.DOTALL | re.IGNORECASE
    )

    @staticmethod
    def extract(text: str, source: str = "", target: str = "high_school",
                difficulty: str = "medium") -> list:
        """返回 build_question 后的 dict 列表，每题带 _q_num 字段。"""
        questions = []
        # 统一换行
        text = re.sub(r"\r\n|\r", "\n", text)

        # 1. 按题号切分成块
        blocks = GenericQuestionExtractor._split_blocks(text)

        for q_num, block in blocks:
            q = GenericQuestionExtractor._parse_block(block, source, target, difficulty)
            if q:
                q["_q_num"] = q_num  # 保留原始题号，供答案匹配
                questions.append(q)
        return questions

    @staticmethod
    def _split_blocks(text: str) -> list:
        """按题号切分，返回 (题号, 内容) 列表。"""
        matches = list(GenericQuestionExtractor.Q_SPLIT_RE.finditer(text))
        if not matches:
            return []
        blocks = []
        for i, m in enumerate(matches):
            # 提取题号：第1题 / 1. / (1)
            q_num = m.group(1) or m.group(2) or m.group(3)
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            blocks.append((int(q_num), text[start:end].strip()))
        return blocks

    @staticmethod
    def _parse_block(block: str, source: str, target: str, difficulty: str):
        """解析单个题目块。"""
        # 先切分出选项 A/B/C/D
        parts = GenericQuestionExtractor.OPT_RE.split(block)
        # parts 结构: [stem, 'A', optA, 'B', optB, 'C', optC, 'D', tail]
        if len(parts) < 9:
            return None

        stem = parts[0].strip()
        opts = {}
        for i in range(1, len(parts) - 1, 2):
            label = parts[i]
            content = parts[i + 1]
            opts[label] = content.strip()

        # 至少要有 4 个选项（A/B/C/D），支持 E
        if sum(1 for k in "ABCD" if k in opts) < 4:
            return None

        last_opt = "E" if "E" in opts else "D"
        tail = opts.get(last_opt, "")
        # 从 tail 中分离选项内容和答案/解析
        opt_clean, answer, analysis = GenericQuestionExtractor._split_tail(tail)
        opts[last_opt] = opt_clean

        if len(stem) < 5:
            return None
        return build_question(stem, opts, answer, analysis, source=source,
                              target=target, difficulty=difficulty)

    @staticmethod
    def _split_tail(tail: str) -> tuple:
        """把选项 D 尾部按 答案/解析 切分。"""
        # 先尝试找解析
        ana_m = GenericQuestionExtractor.ANA_RE.search(tail)
        analysis = ana_m.group(1).strip() if ana_m else ""
        # 去除解析后剩余部分
        remainder = tail[:ana_m.start()].strip() if ana_m else tail

        # 从 remainder 找答案
        ans_m = GenericQuestionExtractor.ANS_RE.search(remainder)
        answer = "A"
        if ans_m:
            ans = ans_m.group(1).strip().upper()
            if len(ans) == 1 and ans in "ABCD":
                answer = ans
            elif len(ans) > 1 and all(c in "ABCD" for c in ans):
                answer = {k: (k in ans) for k in "ABCD"}
            remainder = remainder[:ans_m.start()].strip()

        return remainder, answer, analysis


# ═══════════════════════════════════════════════════════════════
# 示例解析器 1：通用 HTML 页面（用 BeautifulSoup 提取文本后走启发式）
# ═══════════════════════════════════════════════════════════════

@register_parser
class GenericHtmlParser(SiteParser):
    name = "generic_html"

    def can_handle(self, url: str) -> bool:
        return True  # 兜底解析器

    def parse_list(self, fetcher: EthicalFetcher, url: str, max_pages: int = 1) -> list:
        # 通用列表页只返回自身
        return [url]

    def parse_detail(self, fetcher: EthicalFetcher, url: str) -> list:
        html = fetcher.get_text(url)
        if not html:
            return []
        soup = BeautifulSoup(html, "lxml")
        # 移除脚本样式
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        return GenericQuestionExtractor.extract(text, source=url,
                                                 target=self.target, difficulty=self.difficulty)


# ═══════════════════════════════════════════════════════════════
# 示例解析器 2：国家中小学智慧教育平台（课程页提取，需按实际结构调整）
# ═══════════════════════════════════════════════════════════════

@register_parser
class SmartEduParser(SiteParser):
    name = "smartedu"

    def can_handle(self, url: str) -> bool:
        return "basic.smartedu.cn" in url or "smartedu.cn" in url

    def parse_list(self, fetcher: EthicalFetcher, url: str, max_pages: int = 1) -> list:
        html = fetcher.get_text(url)
        if not html:
            return []
        soup = BeautifulSoup(html, "lxml")
        links = []
        for a in soup.find_all("a", href=True):
            href = a["href"]
            if "tvrSearch" in href or "course" in href or "detail" in href:
                links.append(urljoin(url, href))
        return list(set(links))

    def parse_detail(self, fetcher: EthicalFetcher, url: str) -> list:
        html = fetcher.get_text(url)
        if not html:
            return []
        soup = BeautifulSoup(html, "lxml")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        return GenericQuestionExtractor.extract(text, source=url,
                                                 target=self.target, difficulty=self.difficulty)


# ═══════════════════════════════════════════════════════════════
# PDF 文本提取公共函数
# ═══════════════════════════════════════════════════════════════

import tempfile


def _extract_docx_text(path: str) -> str:
    """使用 python-docx 提取 .docx 文本。"""
    try:
        from docx import Document
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        log.warning(f"python-docx 读取失败: {e}")
        return ""


def _extract_aspose_text(path: str) -> str:
    """使用 aspose-words 提取 .doc/.docx/.pdf 文本。"""
    try:
        import aspose.words as aw
        doc = aw.Document(path)
        return doc.get_text()
    except Exception as e:
        log.warning(f"aspose-words 读取失败: {e}")
        return ""


def extract_document_text(fetcher: EthicalFetcher, url: str) -> str:
    """下载 PDF/DOC/DOCX 并返回提取的文本。"""
    tmp_dir = tempfile.gettempdir()
    os.makedirs(tmp_dir, exist_ok=True)

    # 从 URL/文件名推断扩展名
    lower_url = url.lower()
    if ".docx" in lower_url:
        ext = ".docx"
    elif ".doc" in lower_url:
        ext = ".doc"
    else:
        ext = ".pdf"

    tmp_path = os.path.join(tmp_dir, f"bio_crawler_{int(time.time()*1000)}_{random.randint(0,9999)}{ext}")
    if not fetcher.download(url, tmp_path):
        return ""

    text = ""
    try:
        if ext in (".docx", ".doc"):
            text = _extract_docx_text(tmp_path)
            if not text:
                text = _extract_aspose_text(tmp_path)
        else:
            try:
                import fitz
                doc = fitz.open(tmp_path)
                for page in doc:
                    text += page.get_text()
                doc.close()
            except Exception:
                text = _extract_aspose_text(tmp_path)
    except Exception as e:
        log.error(f"文档读取失败 {url}: {e}")
    finally:
        try:
            os.remove(tmp_path)
        except:
            pass
    return text


def infer_year_from_url(url: str) -> int:
    m = re.search(r"20\d{2}", url)
    return int(m.group(0)) if m else None


def _extract_answer_key(text: str) -> dict:
    """
    从答案文本中提取题号->答案映射。
    支持格式：
    1. 竖排表格：题号/答案交替行（联赛真题标准格式）
    2. "1.A" "2.B" 内联格式
    3. "1-5 AABCD" 范围格式
    返回 {题号: 'A'/'B'/'C'/'D'/'E' 或多选组合}。
    """
    answers = {}

    # ── 模式1: 竖排表格（联赛真题最常见格式）──
    # 文本结构：题号\n1\n2\n3...\n答案\nA\nB\nC...
    # 或多列：题号\n1\n2\n3\n4\n5\n答案\nA\nB\nC\nD\nE
    # 按段落（空行分隔）解析
    _parse_vertical_table(text, answers)

    # ── 模式2: "1.A" "2.B" 内联格式 ──
    for m in re.finditer(r"(?:^|\n)\s*(\d+)\s*[\.．、]?\s*([A-E])\b", text):
        try:
            idx = int(m.group(1))
            if idx not in answers:  # 竖排优先
                answers[idx] = m.group(2)
        except:
            pass

    # ── 模式3: "1-5 AABCD" 范围格式 ──
    for m in re.finditer(r"(\d+)\s*-\s*(\d+)\s+([A-E]+)", text):
        start, end, ans = int(m.group(1)), int(m.group(2)), m.group(3)
        for i, ch in enumerate(ans):
            if start + i <= end and (start + i) not in answers:
                answers[start + i] = ch

    return answers


def _parse_vertical_table(text: str, answers: dict):
    """
    解析竖排表格格式的答案表。
    典型结构（PDF 提取后）：
        题号
        1  2  3  4  5  6  7  8  9  10
        答案
        删除  D  C  A  删除  B  D  B  AB  D
    或竖排：
        题号
        1
        2
        3
        答案
        A
        B
        C
    """
    # 按 "题号" 和 "答案" 标记分段
    # 找所有 "题号" 和 "答案" 的位置
    q_markers = [(m.start(), m.end()) for m in re.finditer(r"题\s*号", text)]
    a_markers = [(m.start(), m.end()) for m in re.finditer(r"答\s*案", text)]

    if not q_markers or not a_markers:
        return

    for qi, q_marker in enumerate(q_markers):
        # 找到紧跟此"题号"之后的第一个"答案"标记
        a_marker = None
        for am in a_markers:
            if am[0] > q_marker[1]:
                a_marker = am
                break
        if not a_marker:
            continue

        # 题号区域：从"题号"后到"答案"前
        q_section = text[q_marker[1]:a_marker[0]].strip()
        # 答案区域：从"答案"后到下一个"题号"或文本末尾
        next_q = q_markers[qi + 1][0] if qi + 1 < len(q_markers) else len(text)
        a_section = text[a_marker[1]:next_q].strip()

        # 提取题号列表
        q_nums = [int(x) for x in re.findall(r"\d+", q_section)]
        # 提取答案列表（每行或每空格分隔）
        # 先按行分割，每行可能包含多个答案（空格分隔）
        a_tokens = []
        for line in a_section.splitlines():
            line = line.strip()
            if not line:
                continue
            # 按空格/制表符分割
            for tok in re.split(r"\s+", line):
                tok = tok.strip()
                if not tok:
                    continue
                # 清理全角字母
                tok = tok.replace("Ａ","A").replace("Ｂ","B").replace("Ｃ","C").replace("Ｄ","D").replace("Ｅ","E")
                a_tokens.append(tok)

        # 配对
        for i, q_num in enumerate(q_nums):
            if i < len(a_tokens):
                tok = a_tokens[i]
                if tok in ("删除", "DEL", "—", "-"):
                    continue
                # 处理 "B或D" / "B 或D"
                tok = re.sub(r"\s*或\s*", "", tok)
                if re.match(r"^[A-E]+$", tok):
                    answers[q_num] = tok


def _build_answer_summary(answers: dict) -> str:
    """生成简洁答案对照表，用于 analysis 字段。"""
    if not answers:
        return ""
    items = []
    for idx in sorted(answers.keys()):
        items.append(f"{idx}:{answers[idx]}")
    return "标准答案对照：" + ", ".join(items)


def _parse_pdf_questions(question_text: str, answer_text: str, source_url: str,
                         target: str, difficulty: str, detail_url: str = "") -> list:
    """从试题 PDF 和答案 PDF 中提取题目，并尝试匹配答案。"""
    year = infer_year_from_url(detail_url or source_url)
    source = detail_url or source_url
    questions = GenericQuestionExtractor.extract(question_text, source=source,
                                                   target=target, difficulty=difficulty)
    answers = _extract_answer_key(answer_text) if answer_text else {}

    # 清理 aspose 水印
    answer_text_clean = re.sub(
        r"Created with an evaluation copy of Aspose\.Words.*?(?=\n)",
        "", answer_text, flags=re.DOTALL
    ) if answer_text else ""
    answer_text_clean = re.sub(r"\x13.*?\x15", "", answer_text_clean)

    for i, q in enumerate(questions, start=1):
        q["year"] = year
        q["source"] = source
        if year and str(year) not in q["tags"]:
            q["tags"].append(str(year))
        # 用原始题号匹配答案表
        q_num = q.pop("_q_num", i)
        if q_num in answers:
            ans = answers[q_num]
            if isinstance(ans, str) and len(ans) == 1 and ans in "ABCDE":
                q["answer"] = {k: (k == ans) for k in "ABCDE"}
            elif isinstance(ans, str) and len(ans) > 1 and all(c in "ABCDE" for c in ans):
                q["answer"] = {k: (k in ans) for k in "ABCDE"}
            # 生成该题专属解析
            ans_display = ans if isinstance(ans, str) else str(ans)
            q["analysis"] = f"正确答案：{ans_display}"
        elif q.get("analysis"):
            # 如果题干自带解析（如答案：B 解析：...），保留
            pass
        else:
            # 无答案也无解析，标记待补充
            q["analysis"] = "【解析待补充】"
    return questions


# ═══════════════════════════════════════════════════════════════
# 示例解析器 3：教育考试院 PDF 真题
# ═══════════════════════════════════════════════════════════════

@register_parser
class PdfExamParser(SiteParser):
    name = "pdf_exam"

    def can_handle(self, url: str) -> bool:
        return url.lower().endswith(".pdf")

    def parse_list(self, fetcher: EthicalFetcher, url: str, max_pages: int = 1) -> list:
        return [url]

    def parse_detail(self, fetcher: EthicalFetcher, url: str) -> list:
        text = extract_document_text(fetcher, url)
        if not text:
            return []
        year = infer_year_from_url(url)
        questions = GenericQuestionExtractor.extract(text, source=url,
                                                       target=self.target, difficulty=self.difficulty)
        for q in questions:
            q["year"] = year
            if year and str(year) not in q["tags"]:
                q["tags"].append(str(year))
        return questions


# ═══════════════════════════════════════════════════════════════
# 解析器 4：中国动物学会 czs.org.cn 联赛试题下载
# ═══════════════════════════════════════════════════════════════

@register_parser
class CzsParser(SiteParser):
    name = "czs"

    def can_handle(self, url: str) -> bool:
        return "czs.org.cn" in url

    def _is_list_page(self, url: str) -> bool:
        return "/lsstxz/index" in url or url.rstrip("/").endswith("/lsstxz")

    def parse_list(self, fetcher: EthicalFetcher, url: str, max_pages: int = 1) -> list:
        html = fetcher.get_rendered_html(url, wait_seconds=2.0)
        if not html:
            return []
        soup = BeautifulSoup(html, "lxml")
        links = []
        for a in soup.find_all("a", href=True):
            href = a["href"]
            if "/lsstxz/art/" in href and href.endswith(".html"):
                links.append(urljoin(url, href))
        return list(set(links))

    def parse_detail(self, fetcher: EthicalFetcher, url: str) -> list:
        html = fetcher.get_rendered_html(url, wait_seconds=2.0)
        if not html:
            return []
        soup = BeautifulSoup(html, "lxml")

        # 找到试题 PDF 和标准答案 PDF
        question_pdf = None
        answer_pdf = None
        for a in soup.find_all("a", href=True):
            href = a["href"]
            text = a.get_text(strip=True)
            full = urljoin(url, href)
            if ".pdf" not in full.lower() and "download?fileUrl=" not in full:
                continue
            if "答案" in text or "标准答案" in text or "answer" in text.lower():
                answer_pdf = full
            elif "试题" in text or "试卷" in text or "question" in text.lower():
                question_pdf = full

        if not question_pdf:
            log.warning(f"czs 详情页未找到试题 PDF: {url}")
            return []

        log.info(f"czs 试题 PDF: {question_pdf}")
        if answer_pdf:
            log.info(f"czs 答案 PDF: {answer_pdf}")

        # 如果没有独立答案 PDF，尝试从 HTML 页面提取内嵌答案表
        inline_answer_text = ""
        if not answer_pdf:
            page_text = soup.get_text(separator="\n", strip=True)
            if "题号" in page_text and "答案" in page_text:
                # 截取“答案”附近文本作为答案来源
                idx = page_text.find("答案")
                if idx != -1:
                    inline_answer_text = page_text[idx:idx + 4000]
                    log.info("czs 从 HTML 提取内嵌答案表")

        q_text = extract_document_text(fetcher, question_pdf)
        a_text = extract_document_text(fetcher, answer_pdf) if answer_pdf else inline_answer_text
        return _parse_pdf_questions(q_text, a_text, question_pdf, self.target, self.difficulty, detail_url=url)


# ═══════════════════════════════════════════════════════════════
# 解析器 5：中国植物学会 botany.org.cn 生物学联赛
# ═══════════════════════════════════════════════════════════════

@register_parser
class BotanyParser(SiteParser):
    name = "botany"

    def can_handle(self, url: str) -> bool:
        return "botany.org.cn" in url

    def _is_list_page(self, url: str) -> bool:
        return "/swxls/index" in url or url.rstrip("/").endswith("/swxls")

    def parse_list(self, fetcher: EthicalFetcher, url: str, max_pages: int = 1) -> list:
        html = fetcher.get_rendered_html(url, wait_seconds=2.0)
        if not html:
            return []
        soup = BeautifulSoup(html, "lxml")
        links = []
        for a in soup.find_all("a", href=True):
            href = a["href"]
            text = a.get_text(strip=True)
            # 详情页链接：相对 ./YYYYMM/tXXXX.html 或绝对 /swxls/YYYYMM/tXXXX.html
            full = urljoin(url, href)
            if re.search(r"/swxls/\d{6}/t\d+[_-]?\d+\.html", full) and ("试题" in text or "答案" in text):
                links.append(full)
        # 如果 max_pages > 1，继续翻页
        if max_pages > 1:
            base = url.rstrip(".html")
            for p in range(1, max_pages):
                page_url = f"{base}_{p}.html" if p > 1 else f"{base}.html"
                if page_url == url:
                    continue
                html2 = fetcher.get_rendered_html(page_url, wait_seconds=2.0)
                if not html2:
                    break
                soup2 = BeautifulSoup(html2, "lxml")
                found = False
                for a in soup2.find_all("a", href=True):
                    href = a["href"]
                    text = a.get_text(strip=True)
                    full = urljoin(page_url, href)
                    if re.search(r"/swxls/\d{6}/t\d+[_-]?\d+\.html", full) and ("试题" in text or "答案" in text):
                        links.append(full)
                        found = True
                if not found:
                    break
        return list(set(links))

    def parse_detail(self, fetcher: EthicalFetcher, url: str) -> list:
        html = fetcher.get_rendered_html(url, wait_seconds=2.0)
        if not html:
            return []
        soup = BeautifulSoup(html, "lxml")

        question_pdf = None
        answer_pdf = None
        for a in soup.find_all("a", href=True):
            href = a["href"]
            text = a.get_text(strip=True)
            full = urljoin(url, href)
            if not full.lower().endswith(".pdf"):
                continue
            if "答案" in text or "answer" in text.lower():
                answer_pdf = full
            elif "试题" in text or "试卷" in text:
                question_pdf = full

        if not question_pdf:
            log.warning(f"botany 详情页未找到试题 PDF: {url}")
            return []

        log.info(f"botany 试题 PDF: {question_pdf}")
        if answer_pdf:
            log.info(f"botany 答案 PDF: {answer_pdf}")

        q_text = extract_document_text(fetcher, question_pdf)
        a_text = extract_document_text(fetcher, answer_pdf) if answer_pdf else ""
        return _parse_pdf_questions(q_text, a_text, question_pdf, self.target, self.difficulty, detail_url=url)


# ═══════════════════════════════════════════════════════════════
# 主控逻辑
# ═══════════════════════════════════════════════════════════════

class BioCrawler:
    def __init__(self, rate: float = 2.0, target: str = "high_school",
                 difficulty: str = "medium"):
        self.fetcher = EthicalFetcher(rate=rate)
        self.target = target
        self.difficulty = difficulty

    def select_parser(self, url: str) -> SiteParser:
        # 优先使用非通用解析器
        specific = None
        generic = None
        for parser in PARSERS:
            if parser.can_handle(url):
                if parser.name == "generic_html":
                    generic = parser
                else:
                    specific = parser
                    break
        chosen = specific if specific else (generic if generic else GenericHtmlParser())
        chosen.target = self.target
        chosen.difficulty = self.difficulty
        return chosen

    def crawl_html(self, start_urls: list, max_pages: int = 1, max_questions: int = None) -> list:
        results = []
        seen_urls = set()
        queue = list(start_urls)
        pages = 0

        while queue and (max_pages is None or pages < max_pages):
            url = queue.pop(0)
            if url in seen_urls:
                continue
            seen_urls.add(url)

            parser = self.select_parser(url)
            log.info(f"[{parser.name}] 解析: {url}")

            if parser.name in ("generic_html", "smartedu", "czs", "botany"):
                detail_urls = parser.parse_list(self.fetcher, url, max_pages=max_pages)
                for durl in detail_urls:
                    if durl in seen_urls:
                        continue
                    seen_urls.add(durl)
                    qs = parser.parse_detail(self.fetcher, durl)
                    results.extend(qs)
                    log.info(f"  详情页 {durl} 提取 {len(qs)} 题")
                    if max_questions and len(results) >= max_questions:
                        return results[:max_questions]
            else:
                qs = parser.parse_detail(self.fetcher, url)
                results.extend(qs)
                log.info(f"  提取 {len(qs)} 题")

            pages += 1
            if max_questions and len(results) >= max_questions:
                return results[:max_questions]

        return results

    def crawl_pdfs(self, pdf_urls: list) -> list:
        results = []
        parser = PdfExamParser()
        parser.target = self.target
        parser.difficulty = self.difficulty
        for url in pdf_urls:
            qs = parser.parse_detail(self.fetcher, url)
            results.extend(qs)
            log.info(f"[PDF] {url} 提取 {len(qs)} 题")
        return results


def sync_to_supabase(questions: list, batch_size: int = 50) -> bool:
    """把爬取的题目批量同步到 Supabase（复用 server.py 的 sb_upsert_batch）。"""
    if not HAS_SERVER:
        log.error("未导入 server.py，无法同步到 Supabase")
        return False
    if not questions:
        log.warning("没有题目可同步")
        return False

    total = len(questions)
    success = 0
    for i in range(0, total, batch_size):
        batch = questions[i:i + batch_size]
        if sb_upsert_batch(batch):
            success += len(batch)
            log.info(f"Supabase 同步: {min(i + batch_size, total)}/{total}")
        else:
            log.error(f"Supabase 同步失败: 批次 {i}-{i + batch_size}")
    log.info(f"Supabase 同步完成: {success}/{total}")
    return success == total


def merge_into_pool(crawled_path: str, pool_path: str):
    """把爬取的题目合并进 pool.json，按 id 去重。"""
    crawled = json.loads(Path(crawled_path).read_text("utf-8"))
    pool = []
    if Path(pool_path).exists():
        try:
            pool = json.loads(Path(pool_path).read_text("utf-8"))
        except Exception as e:
            log.error(f"读取 {pool_path} 失败: {e}")
            return

    existing_ids = {q.get("id") for q in pool}
    added = 0
    for q in crawled:
        if q.get("id") not in existing_ids:
            pool.append(q)
            existing_ids.add(q["id"])
            added += 1

    Path(pool_path).write_text(json.dumps(pool, ensure_ascii=False, indent=2), "utf-8")
    log.info(f"合并完成: 新增 {added} 题，pool 共 {len(pool)} 题")


def main():
    parser = argparse.ArgumentParser(description="BioQuest 高中生物题目爬虫")
    parser.add_argument("--mode", choices=["html", "pdf", "merge"], default="html",
                        help="爬取模式: html=网页, pdf=PDF真题, merge=合并到pool.json")
    parser.add_argument("--start-urls", type=str, default="",
                        help="起始URL，多个用逗号分隔")
    parser.add_argument("--pdf-urls", type=str, default="",
                        help="PDF URL，多个用逗号分隔")
    parser.add_argument("--max-pages", type=int, default=1,
                        help="HTML模式最大爬取页数")
    parser.add_argument("--max-questions", type=int, default=None,
                        help="最多提取题目数")
    parser.add_argument("--rate", type=float, default=2.0,
                        help="每秒请求数，默认2")
    parser.add_argument("--target", choices=["high_school", "competition", "both"], default="high_school",
                        help="题目受众: high_school=高考, competition=竞赛, both=混合")
    parser.add_argument("--difficulty", choices=["basic", "medium", "hard", "league", "national"], default="medium",
                        help="题目难度")
    parser.add_argument("--output", type=str, default="data/crawled_questions.json",
                        help="输出JSON路径")
    parser.add_argument("--merge-to", type=str, default=None,
                        help="合并到指定 pool.json")
    parser.add_argument("--sync", action="store_true",
                        help="同步到 Supabase（需 server.py 中的配置可用）")
    args = parser.parse_args()

    if args.mode == "merge":
        merge_into_pool(args.output, args.merge_to or "pool.json")
        return

    crawler = BioCrawler(rate=args.rate, target=args.target, difficulty=args.difficulty)
    results = []

    if args.mode == "html":
        urls = [u.strip() for u in args.start_urls.split(",") if u.strip()]
        if not urls:
            log.error("请提供 --start-urls")
            return
        results = crawler.crawl_html(urls, max_pages=args.max_pages, max_questions=args.max_questions)
    elif args.mode == "pdf":
        urls = [u.strip() for u in args.pdf_urls.split(",") if u.strip()]
        if not urls:
            log.error("请提供 --pdf-urls")
            return
        results = crawler.crawl_pdfs(urls)

    if results:
        out_path = Path(args.output)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(json.dumps(results, ensure_ascii=False, indent=2), "utf-8")
        log.info(f"已保存 {len(results)} 题到 {out_path}")
        if args.merge_to:
            merge_into_pool(str(out_path), args.merge_to)
        if args.sync:
            sync_to_supabase(results)
    else:
        log.warning("未提取到任何题目，请检查目标页面结构或提供自定义解析器。")


if __name__ == "__main__":
    main()
